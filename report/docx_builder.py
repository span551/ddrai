from docx import Document
from docx.shared import Inches
import json
import os


class DDRReportBuilder:

    def __init__(self, merged_path, llm_output_path):
        with open(merged_path) as f:
            self.merged = json.load(f)

        with open(llm_output_path) as f:
            self.llm = json.load(f)

        self.doc = Document()

    def add_heading(self, text):
        self.doc.add_heading(text, level=1)

    def add_paragraph(self, text):
        self.doc.add_paragraph(text)

    # ✅ FIXED IMAGE HANDLER
    def add_images(self, image_list):

        if not image_list:
            self.doc.add_paragraph("Image Not Available")
            return

        for img in image_list:

            print("Trying to load image:", img)  # 🔍 debug

            if os.path.exists(img):
                print("✅ Found:", img)

                try:
                    self.doc.add_picture(img, width=Inches(3))
                except Exception as e:
                    print("❌ Error inserting image:", e)
                    self.doc.add_paragraph("Image Not Available")

            else:
                print("❌ Missing:", img)
                self.doc.add_paragraph("Image Not Available")

    def build(self, output_path):

        # 1️⃣ Property Summary
        self.add_heading("1. Property Issue Summary")
        self.add_paragraph(self.llm.get("property_issue_summary", "Not Available"))

        # 2️⃣ Area-wise Observations
        self.add_heading("2. Area-wise Observations")

        for area in self.merged:
            self.doc.add_heading(area["area"], level=2)

            # Inspection Issues
            for issue in area.get("inspection_issues", []):
                self.add_paragraph(f"- {issue}")

            # Thermal Findings
            for t in area.get("thermal_findings", []):
                self.add_paragraph(
                    f"Thermal: Hotspot {t.get('hotspot', 'NA')}°C, "
                    f"Coldspot {t.get('coldspot', 'NA')}°C "
                    f"(Severity: {t.get('severity', 'NA')})"
                )

            # Conflicts
            for c in area.get("conflicts", []):
                self.add_paragraph(f"⚠ Conflict: {c}")

            # Images
            self.add_paragraph("Inspection Images:")
            self.add_images(area.get("inspection_images", []))

            self.add_paragraph("Thermal Images:")
            self.add_images(area.get("thermal_images", []))

        # 3️⃣ Root Cause
        self.add_heading("3. Probable Root Cause")
        for item in self.llm.get("probable_root_cause", []):
            self.add_paragraph(f"- {item}")

        # 4️⃣ Severity
        self.add_heading("4. Severity Assessment")
        for item in self.llm.get("severity_assessment", []):
            self.add_paragraph(str(item))

        # 5️⃣ Recommendations
        self.add_heading("5. Recommended Actions")
        for item in self.llm.get("recommended_actions", []):
            self.add_paragraph(f"- {item}")

        # 6️⃣ Notes
        self.add_heading("6. Additional Notes")
        self.add_paragraph(self.llm.get("additional_notes", "Not Available"))

        # 7️⃣ Missing Info
        self.add_heading("7. Missing or Unclear Information")
        for item in self.llm.get("missing_or_unclear_information", []):
            self.add_paragraph(f"- {item}")

        # Save file
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)

        print("✅ DOCX Report Generated")
